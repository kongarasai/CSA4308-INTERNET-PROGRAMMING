import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def generate_report():
    doc = docx.Document()

    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Styles Setup
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # ==========================================
    # COVER PAGE
    # ==========================================
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(36)
    title_p.paragraph_format.space_after = Pt(12)
    run_title = title_p.add_run("CivicConnect Web Platform")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(28)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A) # Municipal Navy

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(36)
    run_sub = sub_p.add_run("A Multi-Tier Full-Stack Web Application for Municipal Grievance Redressal\nCourse Assignment Final Technical Report")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(14)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(0x0D, 0x94, 0x88) # Teal

    meta_table = doc.add_table(rows=6, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Course Code & Title:", "CSA4308 – Internet Programming"),
        ("Project Title:", "CivicConnect Municipal Grievance Platform"),
        ("Bloom's Taxonomy Levels:", "L3 (Apply), L4 (Analyze), L6 (Create)"),
        ("Mapped SDGs:", "UN SDG 9 (Infrastructure) & SDG 11 (Sustainable Cities)"),
        ("Technology Constraints:", "HTML5, CSS3, Vanilla JS, Servlets, JSP/JSTL, XML/XSD/XSLT, SOAP/WSDL"),
        ("Submission Date:", "September 2026")
    ]
    for i, (label, val) in enumerate(meta_data):
        r = meta_table.rows[i]
        r.cells[0].paragraphs[0].add_run(label).bold = True
        r.cells[1].paragraphs[0].add_run(val)
        set_cell_background(r.cells[0], "F8FAFC")
        set_cell_background(r.cells[1], "FFFFFF")
        set_cell_margins(r.cells[0], 80, 80, 100, 100)
        set_cell_margins(r.cells[1], 80, 80, 100, 100)

    doc.add_page_break()

    # Read and append TECHNICAL_REPORT.md content
    report_file = os.path.join("docs", "TECHNICAL_REPORT.md")
    if os.path.exists(report_file):
        with open(report_file, 'r', encoding='utf-8') as f:
            lines = f.readlines()

        in_table = False
        table_data = []

        for line in lines:
            line_str = line.strip()
            
            # Heading 1
            if line_str.startswith("# "):
                doc.add_heading(line_str[2:], level=1)
            # Heading 2
            elif line_str.startswith("## "):
                h = doc.add_heading(line_str[3:], level=2)
                h.paragraph_format.space_before = Pt(16)
                h.paragraph_format.space_after = Pt(6)
            # Heading 3
            elif line_str.startswith("### "):
                h = doc.add_heading(line_str[4:], level=3)
                h.paragraph_format.space_before = Pt(12)
                h.paragraph_format.space_after = Pt(4)
            # Table Markdown Line
            elif line_str.startswith("|") and line_str.endswith("|"):
                parts = [p.strip() for p in line_str.split("|")[1:-1]]
                if all(set(p) <= set("-: ") for p in parts):
                    continue # Divider line
                if not in_table:
                    in_table = True
                    table_data = []
                table_data.append(parts)
            else:
                if in_table:
                    # Flush table
                    if table_data:
                        t = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                        t.alignment = WD_TABLE_ALIGNMENT.CENTER
                        for r_idx, row_items in enumerate(table_data):
                            for c_idx, item in enumerate(row_items):
                                cell = t.rows[r_idx].cells[c_idx]
                                cell.text = item
                                set_cell_margins(cell, 80, 80, 100, 100)
                                if r_idx == 0:
                                    set_cell_background(cell, "1E3A8A")
                                    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                                    cell.paragraphs[0].runs[0].font.bold = True
                                else:
                                    bg = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
                                    set_cell_background(cell, bg)
                    in_table = False
                    table_data = []

                if line_str.startswith("* ") or line_str.startswith("- "):
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(line_str[2:])
                elif line_str:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(6)
                    p.add_run(line_str)

    # Append Individual Reflection
    doc.add_page_break()
    refl_file = os.path.join("docs", "INDIVIDUAL_REFLECTION.md")
    if os.path.exists(refl_file):
        with open(refl_file, 'r', encoding='utf-8') as f:
            for line in f:
                line_str = line.strip()
                if line_str.startswith("# "):
                    doc.add_heading(line_str[2:], level=1)
                elif line_str.startswith("### "):
                    doc.add_heading(line_str[4:], level=3)
                elif line_str.startswith("* ") or line_str.startswith("- "):
                    doc.add_paragraph(line_str[2:], style='List Bullet')
                elif line_str:
                    doc.add_paragraph(line_str)

    output_filename = "CivicConnect_Technical_Report.docx"
    doc.save(output_filename)
    print(f"Successfully generated Word Document Report: {output_filename}")

if __name__ == "__main__":
    generate_report()
